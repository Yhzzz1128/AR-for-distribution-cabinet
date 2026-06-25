import json, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

P = {
    'bf': 'Calibri', 'bs': 11, 'ba': 6, 'bl': 1.25,
    'h1s': 16, 'h1c': '#2E74B5', 'h1b': 18, 'h1a': 10,
    'h2s': 13, 'h2c': '#2E74B5', 'h2b': 14, 'h2a': 7,
    'tw': 9360, 'ti': 120, 'hf': 'E8EEF5',
    'ct': 80, 'cb': 80, 'cs': 120, 'ce': 120,
}

def sf(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml('<w:rFonts %s w:ascii="%s" w:hAnsi="%s" w:eastAsia="%s"/>' % (nsdecls('w'), name, name, name))
        rPr.append(rf)
    else:
        for a in ['w:ascii','w:hAnsi','w:eastAsia']: rf.set(qn(a), name)
    run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))

def ss(para, before=0, after=6, ls=1.25):
    pf = para.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = ls

def scm(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    xml = '<w:tcMar %s><w:top w:w="%d" w:type="dxa"/><w:bottom w:w="%d" w:type="dxa"/><w:start w:w="%d" w:type="dxa"/><w:end w:w="%d" w:type="dxa"/></w:tcMar>' % (nsdecls('w'), top, bottom, start, end)
    tcPr.append(parse_xml(xml))

def scw(cell, w):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for ex in list(tcPr):
        if ex.tag == qn('w:tcW'): tcPr.remove(ex)
    tcPr.append(parse_xml('<w:tcW %s w:w="%d" w:type="dxa"/>' % (nsdecls('w'), w)))

def sti(table, indent):
    tbl = table._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w')); tbl.insert(0, tblPr)
    for ex in list(tblPr):
        if ex.tag == qn('w:tblInd'): tblPr.remove(ex)
    tblPr.append(parse_xml('<w:tblInd %s w:w="%d" w:type="dxa"/>' % (nsdecls('w'), indent)))

def stw(table, w):
    tbl = table._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w')); tbl.insert(0, tblPr)
    for ex in list(tblPr):
        if ex.tag == qn('w:tblW'): tblPr.remove(ex)
    tblPr.append(parse_xml('<w:tblW %s w:w="%d" w:type="dxa"/>' % (nsdecls('w'), w)))

def ab(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w')); tbl.insert(0, tblPr)
    b = parse_xml('<w:tblBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="C0C4CC"/><w:left w:val="single" w:sz="4" w:space="0" w:color="C0C4CC"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="C0C4CC"/><w:right w:val="single" w:sz="4" w:space="0" w:color="C0C4CC"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="C0C4CC"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="C0C4CC"/></w:tblBorders>' % nsdecls('w'))
    tblPr.append(b)

def sh(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcPr.append(parse_xml('<w:shd %s w:fill="%s" w:val="clear"/>' % (nsdecls('w'), hex_color)))

# Load data
path = r'D:\Unity\AR-peidiangui\Assets\Resources\SmartQA_Data.json'
with open(path, 'r', encoding='utf-8-sig') as f:
    raw = json.load(f)

# Categories
cats = [
    ('\u72b6\u6001\u6307\u793a\u706f', ['\u7535\u6e90\u6307\u793a','\u5173\u95ed\u4f4d\u6307\u793a','\u5f00\u5230\u4f4d\u6307\u793a','\u5173\u9600\u8fd0\u884c\u6307\u793a','\u5f00\u9600\u8fd0\u884c\u6307\u793a']),
    ('\u63a7\u5236\u6309\u94ae', ['\u5c31\u5730\u5f00\u9600\u6309\u94ae','\u5c31\u5730\u5173\u9600\u6309\u94ae','\u5c31\u5730\u505c\u6b62\u6309\u94ae']),
    ('\u6a21\u5f0f\u9009\u62e9', ['\u5c31\u5730/\u505c/\u8fdc\u7a0b\u6309\u94ae']),
]

emap = {e['name']: e for e in raw}

# Doc
doc = Document()
for sec in doc.sections:
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)

sty = doc.styles['Normal']
sty.font.name = P['bf']; sty.font.size = Pt(P['bs'])
sty.paragraph_format.space_before = Pt(0)
sty.paragraph_format.space_after = Pt(P['ba'])
sty.paragraph_format.line_spacing = P['bl']

for lbl, psize, color, before, after in [
    ('Heading 1', P['h1s'], P['h1c'], P['h1b'], P['h1a']),
    ('Heading 2', P['h2s'], P['h2c'], P['h2b'], P['h2a']),
]:
    hs = doc.styles[lbl]
    hs.font.name = P['bf']; hs.font.size = Pt(psize); hs.font.bold = True
    hs.font.color.rgb = RGBColor.from_string(color.lstrip('#'))
    hs.paragraph_format.space_before = Pt(before)
    hs.paragraph_format.space_after = Pt(after)
    hs.paragraph_format.line_spacing = 1.25
    hs.paragraph_format.keep_with_next = True

# Title
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('\u914d\u7535\u67dc\u63a7\u5236\u9762\u677f SmartQA \u77e5\u8bc6\u5e93')
sf(tr, name=P['bf'], size=26, bold=True, color='#0B2545')
tp.paragraph_format.space_before = Pt(36); tp.paragraph_format.space_after = Pt(4); tp.paragraph_format.line_spacing = 1.15

sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run('AR \u667a\u80fd\u57f9\u8bad\u7cfb\u7edf \u00b7 \u667a\u80fd\u95ee\u7b54\u6a21\u5757\u901f\u67e5\u624b\u518c')
sf(sr, name=P['bf'], size=13, bold=False, color='#555555')
sp.paragraph_format.space_after = Pt(8)

mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mt = '\u77e5\u8bc6\u6761\u76ee\uff1a%d \u9879  |  \u5206\u7c7b\uff1a%d \u7c7b  |  2026\u5e746\u6708' % (len(raw), len(cats))
mr = mp.add_run(mt)
sf(mr, name=P['bf'], size=10, bold=False, color='#888888')
mp.paragraph_format.space_after = Pt(24)

rp = doc.add_paragraph(); rp.paragraph_format.space_before = Pt(0); rp.paragraph_format.space_after = Pt(20)
rp._p.get_or_add_pPr().append(parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="6" w:space="1" w:color="C0C4CC"/></w:pBdr>' % nsdecls('w')))

# Overview
h1 = doc.add_heading('\u6982\u8ff0', level=1)
ov = ('\u672c\u6587\u6863\u4e3a\u914d\u7535\u67dcAR\u667a\u80fd\u57f9\u8bad\u7cfb\u7edf\u4e2d SmartQA\uff08\u667a\u80fd\u95ee\u7b54\uff09\u6a21\u5757\u7684\u914d\u5957\u77e5\u8bc6\u5e93\u3002'
      'SmartQA \u8986\u76d6\u914d\u7535\u67dc\u63a7\u5236\u9762\u677f\u4e0a\u5e38\u89c1\u7684\u6307\u793a\u706f\u3001\u63a7\u5236\u6309\u94ae\u53ca\u6a21\u5f0f\u9009\u62e9\u5f00\u5173\u7684\u529f\u80fd\u8bf4\u660e\u3001'
      '\u6b63\u5e38/\u5f02\u5e38\u72b6\u6001\u5224\u65ad\u4e0e\u64cd\u4f5c\u8981\u70b9\uff0c'
      '\u5171\u6536\u5f55 %d \u6761\u77e5\u8bc6\u6761\u76ee\uff0c\u5206\u4e3a\u72b6\u6001\u6307\u793a\u706f\u3001\u63a7\u5236\u6309\u94ae\u3001\u6a21\u5f0f\u9009\u62e9\u4e09\u5927\u7c7b\u3002'
      '\u5b66\u5458\u53ef\u901a\u8fc7\u81ea\u7136\u8bed\u8a00\u63d0\u95ee\uff0c\u7cfb\u7edf\u81ea\u52a8\u5339\u914d\u5173\u952e\u8bcd\u5e76\u8fd4\u56de\u5bf9\u5e94\u77e5\u8bc6\u70b9\u3002') % len(raw)
p = doc.add_paragraph(ov); ss(p, before=0, after=6, ls=1.25)

# Category index table
h1 = doc.add_heading('\u77e5\u8bc6\u5206\u7c7b\u7d22\u5f15', level=1)
itab = doc.add_table(rows=1, cols=3)
stw(itab, P['tw']); sti(itab, P['ti']); ab(itab)
cw = [int(0.6*1440), int(1.4*1440), int(4.5*1440)]
for ci, w in enumerate(cw):
    for row in itab.rows: scw(row.cells[ci], w)

hdr = itab.rows[0]
for i, lab in enumerate(['\u5e8f\u53f7','\u5206\u7c7b','\u5305\u542b\u6761\u76ee']):
    cell = hdr.cells[i]; sh(cell, P['hf']); scm(cell, P['ct'], P['cb'], P['cs'], P['ce'])
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i==0 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(lab); sf(r, name=P['bf'], size=10, bold=True, color='#2E74B5')

for idx, (cn, names) in enumerate(cats):
    row = itab.add_row()
    titles = '\u3001'.join(names)
    for i, (txt, align, bold) in enumerate([
        (str(idx+1), WD_ALIGN_PARAGRAPH.CENTER, True),
        (cn, WD_ALIGN_PARAGRAPH.LEFT, True),
        (titles, WD_ALIGN_PARAGRAPH.LEFT, False),
    ]):
        cell = row.cells[i]; scm(cell, P['ct'], P['cb'], P['cs'], P['ce']); scw(cell, cw[i])
        p = cell.paragraphs[0]; p.alignment = align
        r = p.add_run(txt)
        sf(r, name=P['bf'], size=10, bold=bold, color='#333333' if i>0 else '#0B2545')

# Panel layout overview
h1 = doc.add_heading('\u63a7\u5236\u9762\u677f\u5e03\u5c40\u6982\u89c8', level=1)
dp = doc.add_paragraph('\u914d\u7535\u67dc\u63a7\u5236\u9762\u677f\u901a\u5e38\u5305\u542b\u4ee5\u4e0b\u4e3b\u8981\u5143\u4ef6\uff1a')
ss(dp, before=0, after=4, ls=1.25)

bullets = [
    '\u4e0a\u65b9\u4e3a\u72b6\u6001\u6307\u793a\u706f\u533a\u57df\uff1a\u7535\u6e90\u6307\u793a\u706f\u3001\u5f00/\u5173\u9600\u8fd0\u884c\u6307\u793a\u706f\u3001\u5f00/\u5173\u5230\u4f4d\u6307\u793a\u706f',
    '\u4e2d\u95f4\u4e3a\u63a7\u5236\u6309\u94ae\u533a\u57df\uff1a\u5c31\u5730\u5f00\u9600\u3001\u5c31\u5730\u5173\u9600\u3001\u5c31\u5730\u505c\u6b62',
    '\u4e0b\u65b9\u6216\u4fa7\u8fb9\u4e3a\u6a21\u5f0f\u9009\u62e9\u5f00\u5173\uff1a\u5c31\u5730/\u505c/\u8fdc\u7a0b',
]
for b in bullets:
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(0); bp.paragraph_format.space_after = Pt(2); bp.paragraph_format.line_spacing = 1.25
    bp.paragraph_format.left_indent = Inches(0.375); bp.paragraph_format.first_line_indent = Inches(-0.187)
    br = bp.add_run('\u2022 ' + b); sf(br, name=P['bf'], size=11, bold=False, color='#333333')

# Detailed Q&A
h1 = doc.add_heading('\u77e5\u8bc6\u6761\u76ee\u8be6\u89e3', level=1)

for cn, names in cats:
    h1 = doc.add_heading(cn, level=1)
    for ename in names:
        entry = emap.get(ename)
        if not entry: continue
        h2 = doc.add_heading(entry['name'], level=2)

        kwp = doc.add_paragraph()
        kwp.paragraph_format.space_before = Pt(0); kwp.paragraph_format.space_after = Pt(4); kwp.paragraph_format.line_spacing = 1.15
        kwl = kwp.add_run('\u5173\u952e\u8bcd\uff1a'); sf(kwl, name=P['bf'], size=9, bold=True, color='#888888')
        kwt = kwp.add_run('\u3001'.join(entry['keywords'])); sf(kwt, name=P['bf'], size=9, bold=False, color='#888888')

        # Parse answer into structured sections
        ans = entry['answer']
        segs = re.split(r'(\u3010[^\u3011]+\u3011)', ans)
        for seg in segs:
            seg = seg.strip()
            if not seg: continue
            if seg.startswith('\u3010') and seg.endswith('\u3011'):
                label = seg[1:-1]
                if label == entry['name']: continue
                shp = doc.add_paragraph()
                shp.paragraph_format.space_before = Pt(6); shp.paragraph_format.space_after = Pt(2)
                shp.paragraph_format.line_spacing = 1.20
                shr = shp.add_run(label); sf(shr, name=P['bf'], size=11, bold=True, color='#1F3A5F')
            else:
                cp = doc.add_paragraph()
                cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(P['ba'])
                cp.paragraph_format.line_spacing = P['bl']
                cr = cp.add_run(seg); sf(cr, name=P['bf'], size=P['bs'], bold=False, color='#333333')

# Footer: page numbers
section = doc.sections[0]
footer = section.footer; footer.is_linked_to_previous = False
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_before = Pt(4); fp.paragraph_format.space_after = Pt(0)

r1 = fp.add_run(); sf(r1, name=P['bf'], size=9, bold=False, color='#888888')
r1._r.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w')))
r2 = fp.add_run(); sf(r2, name=P['bf'], size=9, bold=False, color='#888888')
r2._r.append(parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w')))
r3 = fp.add_run(); sf(r3, name=P['bf'], size=9, bold=False, color='#888888')
r3._r.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w')))

out = r'D:\Unity\AR-peidiangui\SmartQA_Knowledge_Base.docx'
doc.save(out)
print('Saved: ' + out)
print('Entries: %d, Categories: %d' % (len(raw), len(cats)))
