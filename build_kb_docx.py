"""Build a compact reference Word document from the knowledge base JSON."""
import json, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Preset: compact_reference_guide ────────────────────────────────────
PRESET = {
    "base_font": "Calibri",
    "body": {"size": 11, "before": 0, "after": 6, "line": 1.25},
    "h1": {"size": 16, "color": "2E74B5", "before": 18, "after": 10},
    "h2": {"size": 13, "color": "2E74B5", "before": 14, "after": 7},
    "h3": {"size": 12, "color": "1F4D78", "before": 10, "after": 5},
    "table_width_dxa": 9360,
    "table_indent_dxa": 120,
    "cell_margins_dxa": {"top": 80, "bottom": 80, "start": 120, "end": 120},
    "header_fill": "E8EEF5",
    "page_margins": {"top": 1.0, "right": 1.0, "bottom": 1.0, "left": 1.0},  # inches
}

JSON_PATH = r"D:\Unity\AR-peidiangui\Assets\Resources\OperationData.json"
OUTPUT = r"D:\Unity\AR-peidiangui\配电柜知识库快速索引.docx"

# ── Load data ──────────────────────────────────────────────────────────
with open(JSON_PATH, "r", encoding="utf-8") as f:
    data = json.load(f)

# Categorize
cat_map = {}
for item in data:
    cmd = item["command"]
    if cmd in ("停电操作","送电操作","高压停电操作","高压送电操作","紧急停电操作","倒闸操作","低压旁路操作","AR系统操作","五防闭锁"):
        cat_map.setdefault("操作流程", []).append(item)
    elif cmd in ("熔断器故障","断路器跳闸","接触器故障","热继电器故障","电压异常","电流异常","接地故障","母排故障","电缆故障","开关拒动","指示灯不亮","电流表异常","功率因数低","控制回路断线","信号回路异常","PT二次回路"):
        cat_map.setdefault("故障处理", []).append(item)
    elif cmd in ("日常巡检","月度巡检","季度巡检","年度检修","红外测温巡检","柜体外观检查","接地系统巡检","定期清洁","螺栓紧固","绝缘摇测","回路电阻测试","二次回路检查","机械润滑","检修记录管理"):
        cat_map.setdefault("巡检维护", []).append(item)
    elif cmd in ("个人防护装备","验电规范","挂接地线","安全距离","触电急救","操作票制度"):
        cat_map.setdefault("安全规范", []).append(item)
    elif cmd in ("电气火灾","雷击跳闸","洪涝处理","小动物短路","人身触电事故","大面积停电"):
        cat_map.setdefault("应急处置", []).append(item)
    elif cmd in ("备品备件管理","熔断器选型","断路器选型","负载率管理"):
        cat_map.setdefault("设备管理", []).append(item)
    elif cmd in ("蓄电池维护","充电机故障","直流接地故障","直流屏维护"):
        cat_map.setdefault("直流系统", []).append(item)
    elif cmd in ("温升管理","通风散热","谐波治理","三相不平衡","自动无功补偿","UPS维护","剩余电流监测","防雷接地检测"):
        cat_map.setdefault("监测诊断", []).append(item)
    else:
        cat_map.setdefault("其他", []).append(item)

ORDER = ["操作流程","故障处理","巡检维护","安全规范","应急处置","设备管理","直流系统","监测诊断","其他"]
categories = [(cat, cat_map.get(cat, [])) for cat in ORDER if cat_map.get(cat)]

# ── Build document ─────────────────────────────────────────────────────
doc = Document()

# --- Page setup ---
section = doc.sections[0]
m = PRESET["page_margins"]
section.top_margin = Inches(m["top"])
section.bottom_margin = Inches(m["bottom"])
section.left_margin = Inches(m["left"])
section.right_margin = Inches(m["right"])

# --- Set default font ---
style = doc.styles["Normal"]
style.font.name = PRESET["base_font"]
style.font.size = Pt(PRESET["body"]["size"])
style.paragraph_format.space_before = Pt(PRESET["body"]["before"])
style.paragraph_format.space_after = Pt(PRESET["body"]["after"])
style.paragraph_format.line_spacing = PRESET["body"]["line"]
rpr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{PRESET["base_font"]}" w:hAnsi="{PRESET["base_font"]}" w:eastAsia="Microsoft YaHei"/>')
rpr.append(rFonts)

# --- Heading styles ---
for level, h_cfg in [("Heading 1", PRESET["h1"]), ("Heading 2", PRESET["h2"]), ("Heading 3", PRESET["h3"])]:
    hs = doc.styles[level]
    hs.font.name = PRESET["base_font"]
    hs.font.size = Pt(h_cfg["size"])
    hs.font.color.rgb = RGBColor.from_string(h_cfg["color"])
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(h_cfg["before"])
    hs.paragraph_format.space_after = Pt(h_cfg["after"])
    hs.paragraph_format.line_spacing = PRESET["body"]["line"]
    rpr = hs.element.get_or_add_rPr()
    rpr.append(rFonts)

# --- Helper: add table ---
def add_kb_table(doc, items, compact=True):
    """Add a compact 3-column table: 关键词 | 问题标题 | 操作步骤"""
    # Column widths (DXA): sum = 9360
    col_widths = [1690, 2710, 4960]  # ~1.17in + 1.88in + 3.45in = 6.5in
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width and indent
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{PRESET["table_width_dxa"]}" w:type="dxa"/>')
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)
    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="{PRESET["table_indent_dxa"]}" w:type="dxa"/>')
    existing_ind = tblPr.find(qn("w:tblInd"))
    if existing_ind is not None:
        tblPr.remove(existing_ind)
    tblPr.append(tblInd)

    # Set grid column widths
    tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}><w:gridCol w:w="{col_widths[0]}"/><w:gridCol w:w="{col_widths[1]}"/><w:gridCol w:w="{col_widths[2]}"/></w:tblGrid>')
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tbl.insert(0, tblGrid)

    # Cell margins for all cells
    cm_top = PRESET["cell_margins_dxa"]["top"]
    cm_bot = PRESET["cell_margins_dxa"]["bottom"]
    cm_st = PRESET["cell_margins_dxa"]["start"]
    cm_end = PRESET["cell_margins_dxa"]["end"]

    # Header row
    hdr_fill = PRESET["header_fill"]
    for col_idx, header_text in enumerate(["搜索关键词", "显示问题标题", "操作步骤"]):
        cell = table.rows[0].cells[col_idx]
        cell._tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[col_idx]}" w:type="dxa"/>')
        cell._tc.tcPr.append(tcW)
        # Cell margins
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{cm_top}" w:type="dxa"/><w:bottom w:w="{cm_bot}" w:type="dxa"/><w:start w:w="{cm_st}" w:type="dxa"/><w:end w:w="{cm_end}" w:type="dxa"/></w:tcMar>')
        cell._tc.tcPr.append(tcMar)
        # Header fill
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hdr_fill}"/>')
        cell._tc.tcPr.append(shd)
        # Text
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header_text)
        run.font.name = PRESET["base_font"]
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("1F3864")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

    # Data rows
    for item in items:
        row = table.add_row()
        keywords = "、".join(item["keywords"])
        # Compact steps: just 2-3 words per step
        compact_steps = []
        for s in item["steps"]:
            # Keep first 15 chars or until first comma for compactness
            short = s[:20].rstrip("，。,.")
            if len(s) > 20:
                short += "…"
            compact_steps.append(short)
        steps_text = " → ".join(compact_steps[:3])
        if len(item["steps"]) > 3:
            steps_text += f" …（共{len(item['steps'])}步）"

        vals = [keywords, item["title"], steps_text]
        for col_idx, val in enumerate(vals):
            cell = row.cells[col_idx]
            cell._tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[col_idx]}" w:type="dxa"/>')
            cell._tc.tcPr.append(tcW)
            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{cm_top}" w:type="dxa"/><w:bottom w:w="{cm_bot}" w:type="dxa"/><w:start w:w="{cm_st}" w:type="dxa"/><w:end w:w="{cm_end}" w:type="dxa"/></w:tcMar>')
            cell._tc.tcPr.append(tcMar)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.name = PRESET["base_font"]
            run.font.size = Pt(9)
            if col_idx == 0:
                run.font.color.rgb = RGBColor.from_string("1F4D78")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
    return table

# ── Build content ──────────────────────────────────────────────────────
# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("配电柜AR系统 — 知识库快速索引")
title_run.font.name = PRESET["base_font"]
title_run.font.size = Pt(20)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor.from_string("1F3864")
title_p.paragraph_format.space_after = Pt(4)

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run(f"共 {len(data)} 条知识条目 · {len(categories)} 个分类 · 搜索关键词即可定位操作流程")
sub_run.font.name = PRESET["base_font"]
sub_run.font.size = Pt(9)
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
sub_p.paragraph_format.space_after = Pt(14)

# Each category as H1 + table
for cat_name, items in categories:
    doc.add_heading(f"{cat_name}（{len(items)}条）", level=1)
    add_kb_table(doc, items)
    # Small spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)
    spacer_run = spacer.add_run("")
    spacer_run.font.size = Pt(2)

doc.save(OUTPUT)
print(f"OK: {OUTPUT}")
print(f"Entries: {len(data)}, Categories: {len(categories)}")
